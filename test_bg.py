from docx import Document
from docx.shared import Cm
from docx.oxml import parse_xml

def make_image_floating(run, width_cm, height_cm):
    drawing = run._r.xpath('.//w:drawing')[0]
    inline = drawing.xpath('.//wp:inline')[0]
    
    anchor_xml = f'''
    <wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
               xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
               distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0"
               behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="page">
            <wp:align>center</wp:align>
        </wp:positionH>
        <wp:positionV relativeFrom="page">
            <wp:align>center</wp:align>
        </wp:positionV>
        <wp:extent cx="{int(width_cm * 360000)}" cy="{int(height_cm * 360000)}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapNone/>
        <wp:docPr id="999" name="Frame"/>
        <wp:cNvGraphicFramePr/>
    </wp:anchor>
    '''
    anchor = parse_xml(anchor_xml)
    graphic = inline.xpath('.//a:graphic')[0]
    anchor.append(graphic)
    drawing.replace(inline, anchor)

doc = Document()
p = doc.add_paragraph()
r = p.add_run()
r.add_picture(r'images\frame.jpg', width=Cm(18), height=Cm(26))
make_image_floating(r, 18, 26)

doc.add_paragraph('This text should be ON TOP of the image!')

doc.save('scratch/test_bg.docx')
print('Done bg image')
