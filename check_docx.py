# -*- coding: utf-8 -*-
import sys; sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document('d:/chatapp/bao_cao_thanh_vien_2.docx')
NS = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'

for i, p in enumerate(doc.paragraphs):
    has_img = bool(p._element.findall(f'.//{NS}inline'))
    t = p.text.strip()
    if has_img:
        print(f'[{i:3d}] IMAGE  | next caption: ', end='')
        if i+1 < len(doc.paragraphs):
            print(doc.paragraphs[i+1].text.strip()[:80])
        else:
            print('(none)')
    elif 'Hình' in t and len(t) < 100 and ':' in t:
        print(f'[{i:3d}] CAPTION| {t}')

# Also check sections
print('\n--- SECTIONS ---')
for j, sec in enumerate(doc.sections):
    print(f'Section {j}: orient={sec.orientation}, w={sec.page_width}, h={sec.page_height}')
