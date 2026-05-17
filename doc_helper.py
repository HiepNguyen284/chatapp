"""
doc_helper.py — Professional Word document formatting helpers (v4)
Fixes: page-number LEFT, info-box keep-together + LEFT align, reduced spacing
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def insert_zwsp(text):
    """Insert Zero-Width Space (\u200B) to allow MS Word to break long strings gracefully."""
    if not text or not isinstance(text, str):
        return text
    text = re.sub(r'([a-z])([A-Z])', lambda m: m.group(1) + '\u200b' + m.group(2), text)
    text = re.sub(r'([/=?&_.\-()\[\]{}])', lambda m: m.group(1) + '\u200b', text)
    return text

# ── Palette ────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x37, 0x6C)
STEEL  = RGBColor(0x1F, 0x5E, 0x9A)
SLATE  = RGBColor(0x2E, 0x75, 0xB6)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
DARK   = RGBColor(0x1E, 0x1E, 0x1E)
GREY   = RGBColor(0x55, 0x55, 0x55)

ACCENT_HEX  = '2E75B6'
ALT_HEX     = 'DEEAF1'
BOX_BG_HEX  = 'EBF3FB'
BOX_BD_HEX  = '2E75B6'
CODE_BG_HEX = 'F6F6F6'
CODE_BD_HEX = 'C8C8C8'
SEP_HEX     = 'BBBBBB'

FONT      = 'Times New Roman'
FONT_CODE = 'Consolas'
SZ_BODY   = Pt(14)
SZ_H1     = Pt(16)
SZ_H2     = Pt(15)
SZ_H3     = Pt(14)
LS        = 1.25
INDENT    = Cm(1.27)


# ══════════════════════════════════════════════════════════════
#  XML HELPERS
# ══════════════════════════════════════════════════════════════
def _shade_cell(cell, hex_color):
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    shd = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                    f'w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shd)

def _shade_para(p, hex_color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)

def _para_borders(p, sides, color, sz='8', val='single', space='4'):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for s in sides:
        bd = OxmlElement(f'w:{s}')
        bd.set(qn('w:val'), val)
        bd.set(qn('w:sz'), sz)
        bd.set(qn('w:space'), space)
        bd.set(qn('w:color'), color)
        pBdr.append(bd)
    pPr.append(pBdr)

def _heading_underline(p, color_hex, sz='12'):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bd = OxmlElement('w:bottom')
    bd.set(qn('w:val'), 'single')
    bd.set(qn('w:sz'), sz)
    bd.set(qn('w:space'), '1')
    bd.set(qn('w:color'), color_hex)
    pBdr.append(bd)
    pPr.append(pBdr)

def _keep_with_next(p):
    """Prevent page break between this paragraph and the next."""
    pPr = p._p.get_or_add_pPr()
    kwn = OxmlElement('w:keepNext')
    pPr.append(kwn)

def _keep_together(p):
    """Keep all lines of this paragraph on the same page."""
    pPr = p._p.get_or_add_pPr()
    kl = OxmlElement('w:keepLines')
    pPr.append(kl)

def _run_font(run, name=FONT, size=SZ_BODY, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = size
    run.bold      = bold
    run.italic    = italic
    if color:
        run.font.color.rgb = color
    rPr    = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), name)

def _para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              first_indent=INDENT, left=None, right=None,
              sb=Pt(0), sa=Pt(3), ls=LS):
    pf = p.paragraph_format
    pf.alignment         = align
    pf.first_line_indent = first_indent
    if left  is not None: pf.left_indent  = left
    if right is not None: pf.right_indent = right
    pf.space_before = sb
    pf.space_after  = sa
    pf.line_spacing  = ls


# ══════════════════════════════════════════════════════════════
#  DOCUMENT SETUP
# ══════════════════════════════════════════════════════════════
def create_doc():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.0)

        # Page number — RIGHT aligned
        footer = sec.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT          # ← RIGHT
        fp.paragraph_format.first_line_indent = Cm(0)
        r = fp.add_run()
        r.font.name = FONT
        r.font.size = Pt(11)
        b   = OxmlElement('w:fldChar');   b.set(qn('w:fldCharType'), 'begin')
        ins = OxmlElement('w:instrText'); ins.text = ' PAGE '; ins.set(qn('xml:space'), 'preserve')
        e   = OxmlElement('w:fldChar');   e.set(qn('w:fldCharType'), 'end')
        r._element.append(b)
        r._element.append(ins)
        r._element.append(e)

    # ── Normal style ──────────────────────────────────────────
    ns = doc.styles['Normal']
    ns.font.name             = FONT
    ns.font.size             = Pt(14)
    pf = ns.paragraph_format
    pf.alignment             = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent     = INDENT
    pf.line_spacing          = LS
    pf.space_before          = Pt(0)
    pf.space_after           = Pt(2)    # reduced further for space compression

    # ── Heading styles ────────────────────────────────────────
    for name, sz, clr, sb, sa, pbr in [
        ('Heading 1', SZ_H1, NAVY,  Pt(12), Pt(6),  True),   # page break before
        ('Heading 2', SZ_H2, STEEL, Pt(8), Pt(4),  False),
        ('Heading 3', SZ_H3, SLATE, Pt(6),  Pt(2),  False),
    ]:
        hs = doc.styles[name]
        hs.font.name      = FONT
        hs.font.size      = sz
        hs.font.bold      = True
        hs.font.color.rgb = clr
        hpf = hs.paragraph_format
        hpf.first_line_indent  = Cm(0)
        hpf.alignment          = WD_ALIGN_PARAGRAPH.LEFT
        hpf.space_before       = sb
        hpf.space_after        = sa
        hpf.line_spacing       = LS
        hpf.page_break_before  = pbr   # H1 always starts new page
        hpf.keep_with_next     = True

    # ── List Bullet ───────────────────────────────────────────
    lb = doc.styles['List Bullet']
    lb.font.name = FONT
    lb.font.size = Pt(14)
    lbpf = lb.paragraph_format
    lbpf.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
    lbpf.first_line_indent = Cm(0)
    lbpf.left_indent       = Cm(1.5)
    lbpf.line_spacing      = LS
    lbpf.space_before      = Pt(1)
    lbpf.space_after       = Pt(2)     # tight bullet spacing

    try:
        doc.styles['List Bullet'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    except KeyError:
        pass
        
    return doc


# ══════════════════════════════════════════════════════════════
#  BUILDING BLOCKS
# ══════════════════════════════════════════════════════════════

def title_page(doc, ch_num, ch_title):
    for _ in range(4):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after       = Pt(0)
        sp.paragraph_format.first_line_indent = Cm(0)

    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.first_line_indent = Cm(0)
    p1.paragraph_format.space_after       = Pt(4)
    _run_font(p1.add_run(ch_num), size=Pt(26), bold=True, color=NAVY)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.paragraph_format.first_line_indent = Cm(0)
    rule.paragraph_format.space_before      = Pt(0)
    rule.paragraph_format.space_after       = Pt(10)
    _para_borders(rule, ['bottom'], ACCENT_HEX, sz='18')

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.first_line_indent = Cm(0)
    p2.paragraph_format.space_before      = Pt(10)
    p2.paragraph_format.space_after       = Pt(60)
    _run_font(p2.add_run(ch_title), size=Pt(18), bold=True, color=STEEL)

    doc.add_page_break()


def heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    cfg = {
        1: (SZ_H1, NAVY,  '1A376C', '14'),
        2: (SZ_H2, STEEL, 'AECCE8', '6'),
        3: (SZ_H3, SLATE, None,     None),
    }
    sz, clr, brd_clr, brd_sz = cfg[level]
    for r in h.runs:
        _run_font(r, size=sz, bold=True, color=clr)
    if brd_clr:
        _heading_underline(h, brd_clr, brd_sz)
    return h


def para(doc, text, bold=False, italic=False, indent=True,
         sa=Pt(4), align=WD_ALIGN_PARAGRAPH.JUSTIFY, color=None):
    p = doc.add_paragraph()
    _para_fmt(p, align=align,
              first_indent=INDENT if indent else Cm(0), sa=sa)
    _run_font(p.add_run(insert_zwsp(text)), bold=bold, italic=italic, color=color)
    return p


def para_mixed(doc, parts, indent=True, sa=Pt(4)):
    p = doc.add_paragraph()
    _para_fmt(p, first_indent=INDENT if indent else Cm(0), sa=sa)
    for item in parts:
        txt  = item[0]
        bold = item[1] if len(item) > 1 else False
        ital = item[2] if len(item) > 2 else False
        clr  = item[3] if len(item) > 3 else None
        _run_font(p.add_run(insert_zwsp(txt)), bold=bold, italic=ital, color=clr)
    return p


def bullet(doc, body, prefix=None, indent_cm=1.5):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    _para_fmt(p, first_indent=Cm(0), left=Cm(indent_cm), align=WD_ALIGN_PARAGRAPH.LEFT,
              sb=Pt(1), sa=Pt(2))
    if prefix:
        _run_font(p.add_run(insert_zwsp(prefix + ': ')), bold=True)
    _run_font(p.add_run(insert_zwsp(body)))
    return p


def _tbl_borders(tbl, color_hex, sz='12'):
    """Set all outer borders of a table to solid single line."""
    tblEl = tbl._tbl
    tblPr = tblEl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tblEl.insert(0, tblPr)
    tblBdr = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        bd = OxmlElement(f'w:{edge}')
        bd.set(qn('w:val'),   'single')
        bd.set(qn('w:sz'),    sz)
        bd.set(qn('w:color'), color_hex)
        bd.set(qn('w:space'), '0')
        tblBdr.append(bd)
    # Remove inner grid lines
    for edge in ['insideH', 'insideV']:
        bd = OxmlElement(f'w:{edge}')
        bd.set(qn('w:val'), 'none')
        tblBdr.append(bd)
    tblPr.append(tblBdr)


def _tbl_borders_elegant(tbl, color_hex, sz='4'):
    """Set solid elegant borders on all edges (top, bottom, left, right, insideH, insideV)."""
    tblEl = tbl._tbl
    tblPr = tblEl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tblEl.insert(0, tblPr)
    tblBdr = OxmlElement('w:tblBorders')
    for edge in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
        bd = OxmlElement(f'w:{edge}')
        bd.set(qn('w:val'),   'single')
        bd.set(qn('w:sz'),    sz)
        bd.set(qn('w:color'), color_hex)
        bd.set(qn('w:space'), '0')
        tblBdr.append(bd)
    tblPr.append(tblBdr)


def _cell_margins(cell, top=72, bottom=72, left=113, right=113):
    """Set inner cell padding (values in twentieths of a point)."""
    tcPr  = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'),    str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)


def info_box(doc, title, *body_lines):
    """
    Table-based shaded box with a solid blue border.
    Using a table ensures the border is always a single solid line —
    no paragraph-border stitching artefacts.
    title      — bold header line
    *body_lines — each becomes its own LEFT-aligned paragraph inside the cell
    """
    LEFT = WD_ALIGN_PARAGRAPH.LEFT

    # 1×1 table
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Remove default Table Grid style borders and apply custom
    _tbl_borders(tbl, BOX_BD_HEX, sz='14')

    cell = tbl.rows[0].cells[0]
    _shade_cell(cell, BOX_BG_HEX)
    _cell_margins(cell, top=80, bottom=80, left=130, right=130)

    # Prevent the row from splitting across pages
    tr    = tbl.rows[0]._tr
    trPr  = tr.get_or_add_trPr()
    cant  = OxmlElement('w:cantSplit')
    trPr.append(cant)

    # ── Title paragraph (use the cell's built-in first paragraph) ──
    p_t = cell.paragraphs[0]
    p_t.alignment                     = LEFT
    p_t.paragraph_format.alignment    = LEFT
    p_t.paragraph_format.first_line_indent = Cm(0)
    p_t.paragraph_format.space_before = Pt(2)
    p_t.paragraph_format.space_after  = Pt(4)
    p_t.paragraph_format.line_spacing = LS
    _run_font(p_t.add_run(title), bold=True, color=STEEL)

    # ── Body paragraphs ───────────────────────────────────────────
    for i, line in enumerate(body_lines):
        is_last = (i == len(body_lines) - 1)
        p = cell.add_paragraph()
        p.alignment                     = LEFT
        p.paragraph_format.alignment    = LEFT
        p.paragraph_format.first_line_indent = INDENT if i == 0 else Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(6) if is_last else Pt(2)
        p.paragraph_format.line_spacing = LS
        _run_font(p.add_run(line))

    # Space after box
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before      = Pt(0)
    sp.paragraph_format.space_after       = Pt(6)
    sp.paragraph_format.first_line_indent = Cm(0)
    return tbl


def _REMOVED_OLD_MAKE_BOX_PARA(*a, **kw):
    """placeholder — replaced by table-based info_box"""
    pass


def _old_unused_code(doc, title, *body_lines):
    """old paragraph-based info_box — kept for reference only, not called"""
    LEFT = WD_ALIGN_PARAGRAPH.LEFT
    n = len(body_lines)

    def _make_box_para(sides, sb, sa, is_last):
        p = doc.add_paragraph()
        # Force LEFT — never inherit JUSTIFY from Normal
        p.alignment                     = LEFT
        pf = p.paragraph_format
        pf.alignment                    = LEFT
        pf.first_line_indent            = Cm(0)
        pf.left_indent                  = Cm(0.5)
        pf.right_indent                 = Cm(0.5)
        pf.space_before                 = sb
        pf.space_after                  = sa
        pf.line_spacing                 = LS
        _shade_para(p, BOX_BG_HEX)
        _para_borders(p, sides, BOX_BD_HEX, sz='10', space='5')
        _keep_together(p)
        if not is_last:
            _keep_with_next(p)
        return p

    # Title row
    sides_top = ['top', 'left', 'right']
    pt = _make_box_para(sides_top, sb=Pt(5), sa=Pt(3), is_last=(n == 0))
    _run_font(pt.add_run('  ' + title), bold=True, color=STEEL)

    # Body rows
    for i, line in enumerate(body_lines):
        is_last = (i == n - 1)
        sides   = ['bottom', 'left', 'right'] if is_last else ['left', 'right']
        sa_val  = Pt(7) if is_last else Pt(0)
        pb = _make_box_para(sides, sb=Pt(1), sa=sa_val, is_last=is_last)
        indent = Cm(1.27) if i == 0 else Cm(0.5)   # first body line indented
        pb.paragraph_format.first_line_indent = indent
        _run_font(pb.add_run(line))


def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _tbl_borders_elegant(t, 'AECCE8', sz='4')

    for row in t.rows:
        for cell in row.cells:
            _cell_margins(cell, top=100, bottom=100, left=150, right=150)

    if col_widths:
        for ci, w in enumerate(col_widths):
            for cell in t.columns[ci].cells:
                cell.width = Cm(w)

    # Header
    for i, h in enumerate(headers):
        c  = t.rows[0].cells[i]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _shade_cell(c, ACCENT_HEX)
        cp = c.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.first_line_indent = Cm(0)
        cp.paragraph_format.space_before      = Pt(5)
        cp.paragraph_format.space_after       = Pt(5)
        _run_font(cp.add_run(insert_zwsp(h)), bold=True, color=WHITE)

    # Rows
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            if j >= len(headers):
                continue
            c  = t.rows[i + 1].cells[j]
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i % 2 == 1:
                _shade_cell(c, ALT_HEX)
            cp = c.paragraphs[0]
            cp.paragraph_format.first_line_indent = Cm(0)
            cp.paragraph_format.space_before      = Pt(3)
            cp.paragraph_format.space_after       = Pt(3)
            if j == 0:
                cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _run_font(cp.add_run(insert_zwsp(val)), bold=True)
            else:
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _run_font(cp.add_run(insert_zwsp(val)))

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after       = Pt(4)
    sp.paragraph_format.first_line_indent = Cm(0)
    return t


def code_block(doc, *lines):
    """
    LEFT-aligned code block — each argument is one line.
    Lines are joined with w:br so no justify-stretch ever occurs.
    """
    p = doc.add_paragraph()
    p.alignment                         = WD_ALIGN_PARAGRAPH.LEFT   # critical
    pf = p.paragraph_format
    pf.alignment                        = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent                = Cm(0)
    pf.left_indent                      = Cm(1.0)
    pf.right_indent                     = Cm(0.5)
    pf.space_before                     = Pt(4)
    pf.space_after                      = Pt(6)
    pf.line_spacing                     = 1.2
    _shade_para(p, CODE_BG_HEX)
    _para_borders(p, ['top', 'bottom', 'left', 'right'], CODE_BD_HEX, sz='4', space='5')

    for idx, line in enumerate(lines):
        if idx > 0:
            br = OxmlElement('w:br')
            p._p.append(br)
        r = p.add_run(line)
        r.font.name      = FONT_CODE
        r.font.size      = Pt(10)
        r.font.color.rgb = DARK
        r.bold           = False
        r.italic         = False

    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before     = Pt(2)
    p.paragraph_format.space_after      = Pt(10)
    _run_font(p.add_run(text), italic=True, size=Pt(11), color=GREY)


def table_caption(doc, table_id, description):
    """Thêm chú thích dưới bảng: Bảng x.x: Mô tả"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    _run_font(p.add_run(table_id + ': '), bold=True, italic=True, size=Pt(11), color=DARK)
    _run_font(p.add_run(description), italic=True, size=Pt(11), color=DARK)
    return p


def figure_caption(doc, fig_id, description):
    """Thêm chú thích dưới hình ảnh: Hình x.x: Mô tả"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    _run_font(p.add_run(fig_id + ': '), bold=True, italic=True, size=Pt(11), color=DARK)
    _run_font(p.add_run(description), italic=True, size=Pt(11), color=DARK)
    return p


def add_image(doc, image_path, width_cm=10, left_indent_cm=0):
    """Chèn hình ảnh vào document, căn giữa."""
    import os
    if not os.path.exists(image_path):
        para(doc, f'[Hình ảnh không tìm thấy: {image_path}]', italic=True, indent=False)
        return None
    doc.add_picture(image_path, width=Cm(width_cm))
    # doc.add_picture() tạo paragraph mới chứa ảnh -> lấy paragraph cuối để format
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    last_p.paragraph_format.first_line_indent = Cm(0)
    if left_indent_cm != 0:
        last_p.paragraph_format.left_indent = Cm(left_indent_cm)
    last_p.paragraph_format.space_before = Pt(6)
    last_p.paragraph_format.space_after = Pt(2)
    return last_p


def separator(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before     = Pt(8)
    p.paragraph_format.space_after      = Pt(8)
    _para_borders(p, ['bottom'], SEP_HEX, sz='6')


def add_toc(doc, title='MỤC LỤC'):
    """Insert a Word auto-updating Table of Contents field."""
    # Title
    h = doc.add_heading(title, level=1)
    for r in h.runs:
        _run_font(r, size=SZ_H1, bold=True, color=NAVY)
    # Enable page break before for TOC heading so it starts on a new page
    h.paragraph_format.page_break_before = True

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    # TOC field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._element.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._element.append(fldChar2)

    run3 = p.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run3._element.append(fldChar3)


def _make_image_floating(run, width_cm, height_cm):
    from docx.oxml import parse_xml
    drawing = run._r.xpath('.//w:drawing')[0]
    inline = drawing.xpath('.//wp:inline')[0]
    
    anchor_xml = f'''
    <wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
               xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
               distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0"
               behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="page">
            <wp:align>center</wp:align>
        </wp:positionH>
        <wp:positionV relativeFrom="page">
            <wp:align>center</wp:align>
        </wp:positionV>
        <wp:extent cx="{int(width_cm * 360000)}" cy="{int(height_cm * 360000)}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapNone/>
        <wp:docPr id="999" name="Frame"/>
        <wp:cNvGraphicFramePr/>
    </wp:anchor>
    '''
    anchor = parse_xml(anchor_xml)
    graphic = inline.xpath('.//a:graphic')[0]
    anchor.append(graphic)
    drawing.replace(inline, anchor)

def create_cover_page(doc, author_name, author_id):
    """
    Tạo trang bìa giống y hệt (100%) mẫu hình ảnh yêu cầu (Sử dụng frame hoa văn gốc).
    """
    import os
    from docx.enum.table import WD_TABLE_ALIGNMENT

    # Bật tính năng Title Page cho section
    sectPr = doc.sections[0]._sectPr
    titlePg = sectPr.find(qn('w:titlePg'))
    if titlePg is None:
        titlePg = OxmlElement('w:titlePg')
        sectPr.append(titlePg)

    # Chèn Frame (Viền hoa văn) chìm dưới nền
    frame_path = r'd:\chatapp\images\frame.jpg'
    if os.path.exists(frame_path):
        p_frame = doc.paragraphs[0] if len(doc.paragraphs) > 0 else doc.add_paragraph()
        p_frame.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Để không chiếm không gian dòng của chữ, ta set space=0
        p_frame.paragraph_format.space_before = Pt(0)
        p_frame.paragraph_format.space_after = Pt(0)
        p_frame.paragraph_format.line_spacing = 1.0
        r_frame = p_frame.add_run()
        # Khung viền cho trang A4 chuẩn (chừa lề an toàn)
        r_frame.add_picture(frame_path, width=Cm(18), height=Cm(26.5))
        _make_image_floating(r_frame, 18, 26.5)
        

    # Header
    if not os.path.exists(frame_path) and len(doc.paragraphs) == 1 and not doc.paragraphs[0].text:
        p1 = doc.paragraphs[0]
    else:
        p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.first_line_indent = Cm(0)
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(2)
    _run_font(p1.add_run('BỘ KHOA HỌC VÀ CÔNG NGHỆ\n'), size=Pt(14))
    _run_font(p1.add_run('HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG'), size=Pt(14), bold=True)

    # Gạch chân ngắn
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.paragraph_format.first_line_indent = Cm(0)
    p_line.paragraph_format.space_before = Pt(0)
    p_line.paragraph_format.space_after = Pt(12)
    _run_font(p_line.add_run('─────── * ───────'), size=Pt(14))

    # Logo
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.first_line_indent = Cm(0)
    p_logo.paragraph_format.space_after = Pt(20)
    logo_path = r'd:\tieuluan_SAD\images\logo_ptit.png'
    if os.path.exists(logo_path):
        r = p_logo.add_run()
        r.add_picture(logo_path, width=Cm(3.5))
    else:
        _run_font(p_logo.add_run('\n[LOGO PTIT]\n'), size=Pt(14))

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.first_line_indent = Cm(0)
    p_title.paragraph_format.space_after = Pt(18)
    p_title.paragraph_format.line_spacing = 1.3
    _run_font(p_title.add_run('BÁO CÁO BÀI TẬP LỚN\n'), size=Pt(16), bold=True)
    _run_font(p_title.add_run('Môn: Phát triển ứng dụng cho các thiết bị di động\n\n'), size=Pt(14), bold=True)
    _run_font(p_title.add_run('ỨNG DỤNG NHẮN TIN TRỰC TUYẾN — CHATAPP'), size=Pt(18), bold=True, color=NAVY)

    # Info table
    data = [
        ('Giảng viên', ': TS. Nguyễn Hoàng Anh'),
        ('Nhóm (QLĐT)', ': 07'),
        ('Nhóm BTL', ': 10'),
        ('Thành viên nhóm', ': Nguyễn Văn Duy – B22DCCN154,'),
        ('', '  Nguyễn Hoàng Hiệp - B22DCCN298,'),
        ('', '  Nguyễn Quang Minh – B22DCCN538,'),
        ('', '  Đặng Hữu Hoàng Quân - B22DCCN658'),
        ('Thành viên thực hiện', f': {author_name} - {author_id}')
    ]

    table = doc.add_table(rows=len(data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        row.cells[0].width = Cm(5.0)
        row.cells[1].width = Cm(10.0)

    for i, (col1, col2) in enumerate(data):
        c1 = table.cell(i, 0)
        c2 = table.cell(i, 1)
        p_left = c1.paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_left.paragraph_format.first_line_indent = Cm(0)
        p_left.paragraph_format.space_after = Pt(6)
        _run_font(p_left.add_run(col1), size=Pt(14), bold=True)
        
        p_right = c2.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_right.paragraph_format.first_line_indent = Cm(0)
        p_right.paragraph_format.space_after = Pt(6)
        _run_font(p_right.add_run(col2), size=Pt(14), bold=True)

    # Footer
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_footer.paragraph_format.first_line_indent = Cm(0)
    p_footer.paragraph_format.space_before = Pt(60)
    _run_font(p_footer.add_run('Hà nội – 05/2026'), size=Pt(14), italic=True)

    doc.add_page_break()

