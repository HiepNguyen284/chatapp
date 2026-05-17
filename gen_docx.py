# -*- coding: utf-8 -*-
"""Convert 4 bao_cao_thanh_vien .md files to .docx using doc_helper.py format.
Font 14pt Times New Roman, justify, images 10cm width."""
import re, sys, os
sys.path.insert(0, r'd:\chatapp')
from doc_helper import *
from doc_helper import _run_font, _para_fmt, _shade_para, _para_borders
from PIL import Image
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.shared import Cm

IMAGES = r'd:\chatapp\images'

def add_landscape_section(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    return section

def add_portrait_section(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    return section

def get_adjusted_width(img_path, default_w, max_h=15.0):
    """Calculate width so that height does not exceed max_h cm."""
    try:
        with Image.open(img_path) as img:
            w_px, h_px = img.size
        ratio = h_px / w_px
        target_h = default_w * ratio
        if target_h > max_h:
            return max_h / ratio
        return default_w
    except Exception:
        return default_w

def get_img(prefix, fig_name):
    """Get image path for a figure."""
    path = os.path.join(IMAGES, f'{prefix}_hinh_{fig_name}.png')
    if os.path.exists(path):
        return path
    return None

def strip_markdown_symbols(text, preserve_space=False):
    """Remove inline markdown symbols while preserving readable text."""
    if text is None:
        return ''
    cleaned = text if preserve_space else text.strip()
    if not cleaned:
        return text if preserve_space else ''

    # Preserve link information without markdown wrappers.
    cleaned = re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', lambda m: m.group(1) or m.group(2), cleaned)
    cleaned = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1 (\2)', cleaned)

    # Remove common inline markdown wrappers.
    for _ in range(4):
        prev = cleaned
        cleaned = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', cleaned)
        cleaned = re.sub(r'___(.+?)___', r'\1', cleaned)
        cleaned = re.sub(r'\*\*(.+?)\*\*', r'\1', cleaned)
        cleaned = re.sub(r'__(.+?)__', r'\1', cleaned)
        cleaned = re.sub(r'~~(.+?)~~', r'\1', cleaned)
        cleaned = re.sub(r'`([^`]+)`', r'\1', cleaned)
        cleaned = re.sub(r'(?<!\*)\*([^*]+)\*(?!\*)', r'\1', cleaned)
        cleaned = re.sub(r'(?<!_)_([^_]+)_(?!_)', r'\1', cleaned)
        if cleaned == prev:
            break

    # Unescape markdown-escaped punctuation and clean leftovers.
    cleaned = re.sub(r'\\([\\`*_{}\[\]()#+\-.!|>])', r'\1', cleaned)
    cleaned = cleaned.replace('**', '').replace('__', '').replace('~~', '').replace('`', '')
    if not preserve_space:
        cleaned = re.sub(r'\s{2,}', ' ', cleaned).strip()
    return cleaned

def _insert_zwsp(text):
    """Insert Zero-Width Space (\u200B) at logical boundaries to allow MS Word to break long strings gracefully."""
    if not text:
        return text
    # Break at CamelCase / PascalCase
    text = re.sub(r'([a-z])([A-Z])', lambda m: m.group(1) + '\u200b' + m.group(2), text)
    # Break at common safe punctuation (excluding period to avoid breaking acronyms or numbers unnecessarily, though it's usually safe)
    text = re.sub(r'([/=?&_.\-()\[\]{}])', lambda m: m.group(1) + '\u200b', text)
    return text

def add_inline_markdown_runs(paragraph, text):
    """Render inline markdown in a paragraph without showing markdown symbols."""
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
            
        is_bold = False
        is_code = False
        raw_text = part
        
        if part.startswith('**') and part.endswith('**'):
            is_bold = True
            raw_text = strip_markdown_symbols(part[2:-2])
        elif part.startswith('`') and part.endswith('`'):
            is_code = True
            raw_text = strip_markdown_symbols(part[1:-1])
        else:
            raw_text = strip_markdown_symbols(part, preserve_space=True)
            
        if not raw_text:
            continue
            
        # Apply the zero-width space trick to ALL text to fix justified alignment issues "triệt để"
        final_text = _insert_zwsp(raw_text)
        
        r = paragraph.add_run(final_text)
        if is_bold:
            _run_font(r, bold=True)
        elif is_code:
            r.font.name = FONT_CODE
            r.font.size = Pt(11)
        else:
            _run_font(r)

def add_rich_paragraph(doc, text, indent=True):
    """Parse inline markdown and render clean text without markdown symbols."""
    p = doc.add_paragraph()
    _para_fmt(p, first_indent=INDENT if indent else Cm(0), sa=Pt(4))
    add_inline_markdown_runs(p, text)
    return p

def parse_table_lines(lines):
    """Parse markdown table lines into headers, rows, and calculate optimal widths."""
    headers = [strip_markdown_symbols(c.strip()) for c in lines[0].strip('|').split('|')]
    rows = []
    for line in lines[2:]:  # skip separator
        row = [strip_markdown_symbols(c.strip()) for c in line.strip('|').split('|')]
        rows.append(row)
    
    # Calculate optimal column widths (in cm)
    # Total width ~ 16cm (A4 is 21cm - 3cm left - 2cm right)
    max_lengths = [len(h) for h in headers]
    for row in rows:
        for i, cell in enumerate(row):
            if i < len(max_lengths):
                max_lengths[i] = max(max_lengths[i], len(cell))
    
    # Cap maximum length to 45 to prevent one long column from squashing others
    capped_lengths = [min(l, 45) for l in max_lengths]
    total_len = sum(capped_lengths)
    if total_len == 0: total_len = 1
    
    # Base proportional widths
    col_widths = [(l / total_len) * 16.0 for l in capped_lengths]
    
    # Minimum widths based on header length to prevent awkward wrapping
    min_widths = [max(1.2, len(h) * 0.22) for h in headers]
    
    # Iteratively enforce minimums and normalize to 16.0cm
    for _ in range(3):
        current_sum = sum(col_widths)
        if current_sum == 0: break
        scale = 16.0 / current_sum
        col_widths = [w * scale for w in col_widths]
        col_widths = [max(w, min_w) for w, min_w in zip(col_widths, min_widths)]
        
    return headers, rows, col_widths

def convert_report(md_file, prefix, output_file):
    """Convert a single .md report to .docx."""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    doc = create_doc()
    
    # Generate cover page based on filename
    author_name = "Sinh viên"
    author_id = "MSSV"
    if 'thanh_vien_1' in md_file:
        author_name, author_id = "Nguyễn Văn Duy", "B22DCCN154"
    elif 'thanh_vien_2' in md_file:
        author_name, author_id = "Nguyễn Hoàng Hiệp", "B22DCCN298"
    elif 'thanh_vien_3' in md_file:
        author_name, author_id = "Nguyễn Quang Minh", "B22DCCN538"
    elif 'thanh_vien_4' in md_file:
        author_name, author_id = "Đặng Hữu Hoàng Quân", "B22DCCN658"
        
    create_cover_page(doc, author_name, author_id)
    
    # State
    i = 0
    in_code = False
    code_buf = []
    code_lang = ''
    in_mermaid = False
    is_landscape_mode = False
    mermaid_count = 0
    skip_div = False
    figure_captions = re.findall(r'[*_]H.nh (\d+\.\d+):', content)
    
    while i < len(lines):
        line = lines[i].rstrip('\r')
        
        # Skip HTML div tags
        if '<div' in line or '</div>' in line:
            i += 1
            continue
        
        # Skip > note lines
        if line.strip().startswith('> *Ghi ch'):
            i += 1
            continue
            
        # Code/Mermaid block
        if line.strip().startswith('```'):
            if in_code or in_mermaid:
                if in_mermaid:
                    in_mermaid = False
                else:
                    if code_buf:
                        code_block(doc, *code_buf)
                    code_buf = []
                    in_code = False
            else:
                lang = line.strip()[3:].strip()
                if lang == 'mermaid':
                    in_mermaid = True
                else:
                    in_code = True
                    code_lang = lang
            i += 1
            continue
        
        if in_code:
            code_buf.append(line)
            i += 1
            continue
        if in_mermaid:
            i += 1
            continue
        
        # Empty lines
        if not line.strip():
            i += 1
            continue
        
        # Horizontal rule
        if line.strip() == '---':
            i += 1
            continue
        
        # Headings
        if line.startswith('# ') and not line.startswith('## '):
            raw_text = line[2:].strip()
            text = strip_markdown_symbols(raw_text)
            # Title sections get H1
            if text == 'MỤC LỤC':
                add_toc(doc, text)
            elif text in ['DANH SÁCH VIẾT TẮT', 'DANH SÁCH HÌNH', 'DANH SÁCH BẢNG']:
                heading(doc, text, 1)
            elif text.startswith('BẢNG PHÂN CÔNG'):
                heading(doc, text, 1)
            elif text.startswith('BỘ GIÁO DỤC') or text.startswith('BÁO CÁO') or text.startswith('ỨNG DỤNG'):
                pass
            elif text.startswith('Chương'):
                heading(doc, text, 1)
            else:
                heading(doc, text, 1)
            i += 1
            continue
        
        if line.startswith('## '):
            raw_text = line[3:].strip()
            text = strip_markdown_symbols(raw_text)
            
            # Any ## heading: if we're still in landscape, go back to portrait
            if is_landscape_mode:
                add_portrait_section(doc)
                is_landscape_mode = False
            
            # Special: ## 2.5 Biểu đồ tuần tự → switch to landscape for the first diagram
            if 'Biểu đồ tuần tự' in text:
                add_landscape_section(doc)
                is_landscape_mode = True
                
            if raw_text.startswith('[') and 'Tên Trường' in raw_text:  # [Tên Trường]
                pass
            elif text.startswith('Môn:') or text.startswith('KHOA CÔNG NGHỆ'):
                pass
            else:
                heading(doc, text, 2)
            i += 1
            continue
        
        if line.startswith('### '):
            text = strip_markdown_symbols(line[4:].strip())
            
            # Toggle landscape per-subsection
            is_seq_heading = 'Biểu đồ tuần tự' in text
            if is_seq_heading and not is_landscape_mode:
                add_landscape_section(doc)
                is_landscape_mode = True
            elif not is_seq_heading and is_landscape_mode:
                add_portrait_section(doc)
                is_landscape_mode = False
            
            if text.startswith('KHOA'):
                pass
            else:
                heading(doc, text, 3)
            i += 1
            continue
        
        if line.startswith('#### '):
            heading(doc, strip_markdown_symbols(line[5:].strip()), 3)
            i += 1
            continue
        
        # Table
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = []
            while i < len(lines) and '|' in lines[i].rstrip('\r'):
                table_lines.append(lines[i].rstrip('\r'))
                i += 1
            headers, rows, col_widths = parse_table_lines(table_lines)
            add_table(doc, headers, rows, col_widths=col_widths)
            continue
        
        # Caption lines *Bảng x.x:* or *Hình x.x:*
        cap_match = re.match(r'^[*_](Bảng \d+\.\d+|Hình \d+\.\d+):\s*(.*?)[*_]$', line.strip())
        if cap_match:
            is_bang = cap_match.group(1).startswith('Bảng')
            if is_bang and 'Bảng 0.' in cap_match.group(1):
                i += 1
                continue
            
            p_cap = table_caption(
                doc,
                strip_markdown_symbols(cap_match.group(1)),
                strip_markdown_symbols(cap_match.group(2))
            )
            
            # Keep table captions with the table that follows them
            if p_cap and is_bang:
                p_cap.paragraph_format.keep_with_next = True
                
            i += 1
            continue
        
        # Bullet points
        if line.startswith('- ') or line.startswith('  - '):
            is_sub = line.startswith('  - ')
            text = line[4:].strip() if is_sub else line[2:].strip()
            
            # Join wrapped lines
            while i + 1 < len(lines):
                next_line = lines[i+1].rstrip('\r')
                if not next_line.strip() or \
                   next_line.startswith('- ') or next_line.startswith('  - ') or \
                   next_line.startswith('#') or next_line.startswith('|') or \
                   next_line.startswith('```') or next_line.startswith('*Hình') or \
                   next_line.startswith('*Bảng') or next_line.startswith('>'):
                    break
                text += ' ' + next_line.strip()
                i += 1
            
            text = re.sub(r'\s+', ' ', text).strip()
            bp = re.match(r'^\*\*(.*?)\*\*:\s*(.*)', text)
            if bp:
                if is_sub:
                    bullet(doc, strip_markdown_symbols(bp.group(2)), prefix=strip_markdown_symbols(bp.group(1)), indent_cm=2.5)
                else:
                    bullet(doc, strip_markdown_symbols(bp.group(2)), prefix=strip_markdown_symbols(bp.group(1)))
            else:
                if is_sub:
                    bullet(doc, strip_markdown_symbols(text), indent_cm=2.5)
                else:
                    bullet(doc, strip_markdown_symbols(text))
            i += 1
            continue
        
        # Numbered list
        nm = re.match(r'^(\d+)\.\s+(.*)', line)
        if nm:
            text = nm.group(2).strip()
            p = doc.add_paragraph()
            _para_fmt(
                p,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                first_indent=Cm(0),
                left=Cm(1.5),
                sb=Pt(1),
                sa=Pt(2)
            )
            _run_font(p.add_run(f'{nm.group(1)}. '), bold=True)
            add_inline_markdown_runs(p, text)
            i += 1
            continue
        
        # Bold-only lines (like **ChatApp**)
        if line.strip().startswith('**') and line.strip().endswith('**') and line.strip().count('**') == 2:
            text = strip_markdown_symbols(line.strip()[2:-2])
            para(doc, text, bold=True, indent=False)
            i += 1
            continue
        
        # Lines starting with ** (key-value)
        if line.strip().startswith('**') and '**' in line.strip()[2:]:
            add_rich_paragraph(doc, line.strip(), indent=False)
            i += 1
            continue
        
        # Image reference
        m_img = re.match(r'^`?\[image:\s*([^\]]+)\]`?$', line.strip(), re.IGNORECASE)
        if m_img:
            img_name = m_img.group(1).strip()
            # Look ahead to find the caption for sequence detection
            caption_text = ""
            for j in range(i+1, min(i+15, len(lines))):
                m_cap = re.search(r'[*_]Hình (\d+\.\d+):\s*(.*?)[*_]', lines[j])
                if m_cap:
                    caption_text = m_cap.group(2)
                    break
            
            img_path = os.path.join(IMAGES, img_name)
            if os.path.exists(img_path):
                is_sequence = 'tuần tự' in caption_text.lower()
                is_usecase_tong_quan = 'use_case_tong_quan' in img_name.lower()
                
                max_h = 23.0 if is_usecase_tong_quan else 15.0
                w = 14 if ('so_o' in img_name.lower() or 'kien_truc' in img_name.lower() or 'bieu_o_lop' in img_name.lower()) else 12
                if 'giao_dien' in img_name.lower() or 'ui' in img_name.lower() or 'ket_qua' in img_name.lower():
                    w = 7
                    max_h = 14.0
                
                if is_sequence:
                    final_w = get_adjusted_width(img_path, 24.0, max_h=16.0)
                else:
                    final_w = get_adjusted_width(img_path, w, max_h=max_h)
                    
                p_img = add_image(doc, img_path, width_cm=final_w)
                if p_img:
                    p_img.paragraph_format.keep_with_next = True
            i += 1
            continue
        
        # Normal paragraph
        text = line.strip()
        if text:
            # Join wrapped lines
            while i + 1 < len(lines):
                next_line = lines[i+1].rstrip('\r')
                if not next_line.strip() or \
                   next_line.startswith('- ') or next_line.startswith('  - ') or \
                   next_line.startswith('#') or next_line.startswith('|') or \
                   next_line.startswith('```') or next_line.startswith('*Hình') or \
                   next_line.startswith('*Bảng') or next_line.startswith('>'):
                    break
                text += ' ' + next_line.strip()
                i += 1
            
            text = re.sub(r'\s+', ' ', text).strip()
            add_rich_paragraph(doc, text)
        i += 1
    
    doc.save(output_file)
    print(f'Saved: {output_file} ({len(doc.paragraphs)} paragraphs)')

# â”€â”€ Convert all 4 reports â”€â”€
reports = [
    ('bao_cao_thanh_vien_1.md', 'tv1', 'bao_cao_thanh_vien_1.docx'),
    ('bao_cao_thanh_vien_2.md', 'tv2', 'bao_cao_thanh_vien_2.docx'),
    ('bao_cao_thanh_vien_3.md', 'tv3', 'bao_cao_thanh_vien_3.docx'),
    ('bao_cao_thanh_vien_4.md', 'tv4', 'bao_cao_thanh_vien_4.docx'),
]

for md, prefix, out in reports:
    md_path = os.path.join(r'd:\chatapp', md)
    out_path = os.path.join(r'd:\chatapp', out)
    print(f'\nConverting {md}...')
    try:
        convert_report(md_path, prefix, out_path)
    except Exception as e:
        print(f'ERROR: {e}')
        import traceback
        traceback.print_exc()

